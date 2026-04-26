import streamlit as st
import pickle
import re
import html

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

# ---- load models ----
model = pickle.load(open("model/model.pkl", "rb"))
tfidf = pickle.load(open("model/tfidf.pkl", "rb"))
skill_db = pickle.load(open("model/skill_db.pkl", "rb"))


# ---- clean text ----
def clean_text(text):
    text = str(text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = text.lower()
    text = re.sub(r'\W', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text


# ---- analyze with tiered skill db ----
def analyze_resume(resume_text):
    cleaned = clean_text(resume_text)
    vec = tfidf.transform([cleaned])

    predicted_role = model.predict(vec)[0]

    # top 3 predictions
    scores = model.decision_function(vec)[0]
    classes = model.classes_
    top3_indices = scores.argsort()[::-1][:3]
    top3 = [(classes[i], round(scores[i], 2)) for i in top3_indices]
    confidence = round((scores[top3_indices[0]] - scores[top3_indices[1]]) * 100, 1)

    # tiered skill matching
    role_skills = skill_db.get(predicted_role, {})
    high = role_skills.get("high", [])
    medium = role_skills.get("medium", [])
    tools = role_skills.get("tools", [])

    def split(skill_list):
        found = [s for s in skill_list if s in cleaned]
        missing = [s for s in skill_list if s not in cleaned]
        return found, missing

    high_found, high_missing = split(high)
    medium_found, medium_missing = split(medium)
    tools_found, tools_missing = split(tools)

    all_found = high_found + medium_found + tools_found
    all_missing = high_missing + medium_missing + tools_missing
    total = len(high) + len(medium) + len(tools)
    match_score = round(len(all_found) / total * 100, 1) if total else 0

    return {
        "role": predicted_role,
        "top3": top3,
        "confidence": confidence,
        "match_score": match_score,
        "high_found": high_found,
        "high_missing": high_missing,
        "medium_found": medium_found,
        "medium_missing": medium_missing,
        "tools_found": tools_found,
        "tools_missing": tools_missing,
        "all_found": all_found,
        "all_missing": all_missing,
    }


def extract_text_from_pdf(uploaded_file):
    if PdfReader is None:
        raise RuntimeError("PDF support requires package: pypdf")

    reader = PdfReader(uploaded_file)
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages).strip()


def extract_text_from_docx(uploaded_file):
    if Document is None:
        raise RuntimeError("DOCX support requires package: python-docx")

    doc = Document(uploaded_file)
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


def extract_text_from_doc(uploaded_file):
    raw = uploaded_file.getvalue()
    decoded = raw.decode("latin-1", errors="ignore")
    cleaned = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", decoded)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    if len(cleaned) < 40:
        raise RuntimeError("Could not reliably extract text from .doc. Please convert to .docx.")

    return cleaned


# ---- page config ----
st.set_page_config(page_title="Resume Skill Gap Analyzer", page_icon="RG", layout="wide")

if "dynamic_bg_on" not in st.session_state:
    st.session_state.dynamic_bg_on = True

if "resume_input_text" not in st.session_state:
    st.session_state.resume_input_text = ""

if "resume_input_source" not in st.session_state:
    st.session_state.resume_input_source = "manual"

toggle_col, title_col = st.columns([1.1, 4.2], gap="medium")
with toggle_col:
    st.markdown("### Dynamic")
    if st.button(
        "Disable Background" if st.session_state.dynamic_bg_on else "Enable Background",
        key="bg_toggle",
        use_container_width=True,
    ):
        st.session_state.dynamic_bg_on = not st.session_state.dynamic_bg_on
    st.caption("Control background motion independently from title and content.")

with title_col:
    st.markdown(
        """
        <section class="title-stage">
            <p class="title-overline">AI RESUME STUDIO</p>
            <h1>Resume Skill Gap Analyzer</h1>
            <p class="title-copy">
                Dynamic role prediction and actionable skill mapping from plain resume text.
            </p>
        </section>
        """,
        unsafe_allow_html=True,
    )

if st.session_state.dynamic_bg_on:
    st.markdown('<div class="dynamic-bg-layer"></div>', unsafe_allow_html=True)

st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;700&family=IBM+Plex+Mono:wght@400;500&display=swap');

    :root {
        --bg-main: #eaf3ff;
        --bg-card: #ffffff;
        --ink: #0b1a2e;
        --ink-muted: #2f465f;
        --brand: #0b6f8b;
        --brand-soft: #cfe8f5;
        --accent: #c2410c;
        --line: #c8d9e8;
    }

    html, body {
        background: var(--bg-main);
    }

    [data-testid="stAppViewContainer"] {
        background: var(--bg-main);
    }

    [data-testid="stHeader"] {
        background: linear-gradient(180deg, #dcecff 0%, rgba(220, 236, 255, 0.72) 55%, rgba(220, 236, 255, 0) 100%) !important;
    }

    [data-testid="stToolbar"] {
        right: 0.7rem;
        top: 0.45rem;
    }

    .stApp {
        background: var(--bg-main);
        color: var(--ink);
        font-family: 'Space Grotesk', sans-serif;
        overflow-x: hidden;
    }

    .main .block-container {
        max-width: 1100px;
        padding-top: 1.3rem;
        padding-bottom: 2rem;
        position: relative;
        z-index: 4;
    }

    p, li, label, .stMarkdown, .stCaption {
        overflow-wrap: anywhere;
        word-break: normal;
    }

    h1, h2, h3, h4 {
        color: var(--ink);
        font-family: 'Space Grotesk', sans-serif;
        letter-spacing: -0.01em;
    }

    .dynamic-bg-layer {
        position: fixed;
        inset: 0;
        z-index: 0;
        pointer-events: none;
        overflow: hidden;
    }

    .dynamic-bg-layer::before,
    .dynamic-bg-layer::after {
        content: "";
        position: absolute;
        width: 58vw;
        height: 58vw;
        border-radius: 46% 54% 60% 40% / 45% 35% 65% 55%;
        filter: blur(2px);
        opacity: 0.62;
        animation: drift 18s ease-in-out infinite;
    }

    .dynamic-bg-layer::before {
        top: -16vw;
        left: -10vw;
        background: radial-gradient(circle at 35% 35%, #86c7ff 0%, #4ea4ea 45%, transparent 70%);
    }

    .dynamic-bg-layer::after {
        right: -14vw;
        bottom: -22vw;
        background: radial-gradient(circle at 65% 45%, #ffc58e 0%, #f58c42 45%, transparent 72%);
        animation-duration: 23s;
        animation-direction: alternate;
    }

    .title-stage {
        background: linear-gradient(130deg, #0b2342 0%, #12446f 52%, #166f98 100%);
        color: #f8fffe;
        border-radius: 22px;
        padding: 1.25rem 1.4rem;
        margin-bottom: 1.15rem;
        box-shadow: 0 20px 38px rgba(15, 39, 71, 0.3);
        border: 1px solid rgba(213, 235, 255, 0.3);
        animation: fadeUp .6s ease-out;
        position: relative;
        z-index: 5;
    }

    .title-stage h1 {
        color: #f8fffe;
        margin: 0;
        line-height: 1.02;
        letter-spacing: -0.03em;
        font-size: clamp(1.55rem, 4.1vw, 3.2rem);
    }

    .title-overline {
        margin: 0 0 .3rem 0;
        font-family: 'IBM Plex Mono', monospace;
        font-size: .78rem;
        letter-spacing: .18em;
        color: #e7f3ff;
    }

    .title-copy {
        margin: .65rem 0 0 0;
        color: #f0f7ff;
        max-width: 68ch;
        font-size: clamp(.9rem, 1.1vw, 1rem);
        line-height: 1.45;
    }

    [data-testid="column"]:first-child .stButton > button {
        min-height: 3.2rem;
        font-size: .92rem;
        border-radius: 14px;
    }

    .card {
        background: var(--bg-card);
        border: 1px solid var(--line);
        border-radius: 16px;
        padding: .95rem 1rem;
        box-shadow: 0 8px 18px rgba(29, 42, 48, 0.07);
        animation: fadeUp .55s ease-out;
    }

    .metric {
        background: #f2f9ff;
        border: 1px solid #cddfee;
        border-radius: 14px;
        padding: .85rem;
        min-height: 96px;
        animation: fadeUp .55s ease-out;
        overflow-wrap: anywhere;
    }

    .metric .label {
        color: var(--ink-muted);
        font-size: clamp(.72rem, .8vw, .82rem);
        text-transform: uppercase;
        letter-spacing: .05em;
    }

    .metric .value {
        font-family: 'IBM Plex Mono', monospace;
        font-size: clamp(1.05rem, 1.6vw, 1.35rem);
        margin-top: .35rem;
        color: var(--ink);
        font-weight: 500;
        line-height: 1.2;
        overflow-wrap: anywhere;
    }

    .metric.metric-role .value {
        font-size: clamp(.95rem, 1.35vw, 1.22rem);
        letter-spacing: .01em;
    }

    .chip {
        display: inline-block;
        padding: .32rem .55rem;
        border-radius: 999px;
        border: 1px solid #a9cfe8;
        background: #e3f3ff;
        color: #083a5e;
        margin: .18rem .24rem .18rem 0;
        font-size: .84rem;
    }

    .chip.miss {
        border-color: #efbca6;
        background: #ffece3;
        color: #7c2d12;
    }

    .rank-item {
        border-bottom: 1px dashed #c2cfdb;
        padding: .45rem 0;
        animation: fadeUp .55s ease-out;
        overflow-wrap: anywhere;
        line-height: 1.35;
    }

    .rank-item:last-child {
        border-bottom: none;
    }

    .small-note {
        color: #1a3652;
        font-size: clamp(.8rem, .95vw, .9rem);
        font-weight: 600;
        line-height: 1.45;
    }

    .stExpander summary {
        line-height: 1.4;
        overflow-wrap: anywhere;
    }

    [data-testid="stExpander"] {
        border: 1px solid #c8d9e8;
        border-radius: 14px;
        background: rgba(255, 255, 255, 0.45);
        margin-bottom: .55rem;
        overflow: hidden;
    }

    [data-testid="stExpander"] details {
        background: transparent;
    }

    [data-testid="stExpander"] summary {
        background: #eff6fd !important;
        color: var(--ink) !important;
        border-radius: 10px;
        padding: .35rem .55rem;
    }

    [data-testid="stExpander"] details[open] > summary {
        background: #dfeefe !important;
        color: #0b1a2e !important;
        border-bottom: 1px solid #c7daeb;
        border-radius: 10px 10px 0 0;
    }

    [data-testid="stExpander"] summary:hover {
        background: #e5f1fd !important;
    }

    .stButton > button {
        border-radius: 12px;
        border: 1px solid #0b6f8b;
        background: linear-gradient(180deg, #1984a4 0%, #0b6f8b 100%);
        color: #f4fffd;
        font-weight: 700;
        letter-spacing: .01em;
        transition: transform .15s ease, box-shadow .15s ease;
    }

    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 6px 14px rgba(11, 111, 139, 0.33);
    }

    [data-testid="stFileUploaderDropzone"] {
        background: #f5faff !important;
        border: 1px dashed #9dbfda !important;
        border-radius: 12px !important;
        padding: .85rem .9rem !important;
    }

    [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] {
        color: var(--ink) !important;
    }

    [data-testid="stFileUploaderDropzone"] small,
    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] div {
        color: var(--ink-muted);
    }

    [data-testid="stFileUploaderDropzone"] button {
        background: linear-gradient(180deg, #1f90b1 0%, #0b6f8b 100%) !important;
        color: #ecf9ff !important;
        border: 1px solid #0b6f8b !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
    }

    [data-testid="stFileUploaderDropzone"] button span,
    [data-testid="stFileUploaderDropzone"] button div,
    [data-testid="stFileUploaderDropzone"] button p {
        color: #ecf9ff !important;
        font-weight: 700 !important;
        letter-spacing: .01em;
    }

    [data-testid="stFileUploaderDropzone"] button:hover {
        filter: brightness(1.05);
    }

    [data-testid="stFileUploaderFile"] {
        background: #e8f4ff !important;
        border: 1px solid #9fc4df !important;
        border-radius: 10px !important;
    }

    [data-testid="stFileUploaderFile"] [data-testid="stFileUploaderFileName"],
    [data-testid="stFileUploaderFile"] span,
    [data-testid="stFileUploaderFile"] small {
        color: #0b1a2e !important;
    }

    [data-testid="stFileUploaderFile"] [data-testid="stFileUploaderFileName"] {
        opacity: 1 !important;
        visibility: visible !important;
        font-weight: 600 !important;
    }

    [data-testid="stFileUploaderDeleteBtn"] {
        background: #d5e9f9 !important;
        border-radius: 999px !important;
    }

    [data-testid="stFileUploaderDeleteBtn"] svg {
        fill: #0a3e63 !important;
    }

    .upload-success {
        margin-top: .5rem;
        background: linear-gradient(135deg, #0d4e63 0%, #0b3848 100%);
        border: 1px solid #2f7e99;
        color: #e8f7ff;
        border-radius: 12px;
        padding: .75rem .9rem;
        font-weight: 600;
        box-shadow: 0 8px 18px rgba(8, 42, 58, 0.22);
    }

    .upload-success .upload-file-name {
        color: #ffffff;
        font-weight: 700;
        text-decoration: underline;
        text-decoration-color: rgba(255, 255, 255, 0.6);
        text-underline-offset: 2px;
    }

    div[data-baseweb="notification"][kind="warning"] {
        background: linear-gradient(135deg, rgba(255, 245, 224, 0.74) 0%, rgba(255, 234, 204, 0.62) 100%) !important;
        border: 1px solid rgba(218, 168, 104, 0.55) !important;
        border-radius: 12px !important;
        box-shadow: 0 6px 14px rgba(144, 96, 41, 0.10) !important;
        backdrop-filter: blur(1.5px) saturate(108%);
    }

    div[data-baseweb="notification"][kind="warning"] p,
    div[data-baseweb="notification"][kind="warning"] div,
    div[data-baseweb="notification"][kind="warning"] span {
        color: #1a3550 !important;
        font-weight: 600 !important;
    }

    [data-testid="stAlert"] {
        border-radius: 12px !important;
    }

    [data-testid="stAlert"] [data-testid="stMarkdownContainer"] p,
    [data-testid="stAlert"] [data-testid="stMarkdownContainer"] div,
    [data-testid="stAlert"] [data-testid="stMarkdownContainer"] span {
        font-weight: 700 !important;
        letter-spacing: 0.01em;
    }

    [data-testid="stAlert"] [data-testid="stMarkdownContainer"] {
        background: transparent !important;
        border: 0 !important;
        box-shadow: none !important;
        padding: 0 !important;
        margin: 0 !important;
    }

    [data-testid="stAlert"]:has([data-testid="stAlertContentWarning"]) {
        background: linear-gradient(135deg, rgba(255, 245, 224, 0.74) 0%, rgba(255, 234, 204, 0.62) 100%) !important;
        border: 1px solid rgba(218, 168, 104, 0.55) !important;
        box-shadow: 0 6px 14px rgba(144, 96, 41, 0.10) !important;
        backdrop-filter: blur(1.5px) saturate(108%);
    }

    [data-testid="stAlert"] [data-testid="stAlertContentWarning"] {
        background: transparent !important;
        border: 0 !important;
        box-shadow: none !important;
        padding: 0 !important;
        margin: 0 !important;
    }

    [data-testid="stAlert"]:has([data-testid="stAlertContentWarning"]) [data-testid="stMarkdownContainer"] p,
    [data-testid="stAlert"] [data-testid="stAlertContentWarning"] [data-testid="stMarkdownContainer"] p,
    [data-testid="stAlert"] [data-testid="stAlertContentWarning"] p,
    [data-testid="stAlert"] [data-testid="stAlertContentWarning"] div,
    [data-testid="stAlert"] [data-testid="stAlertContentWarning"] span {
        color: #1a3550 !important;
        font-weight: 700 !important;
        text-shadow: 0 1px 0 rgba(255, 255, 255, 0.26);
    }

    div[data-baseweb="notification"][kind="success"] {
        background: #134328 !important;
        border: 1px solid #2c8b54 !important;
        border-radius: 12px !important;
    }

    div[data-baseweb="notification"][kind="success"] p,
    div[data-baseweb="notification"][kind="success"] div,
    div[data-baseweb="notification"][kind="success"] span {
        color: #dcffe9 !important;
    }

    div[data-baseweb="notification"][kind="error"] {
        background: #5a1f1f !important;
        border: 1px solid #d46868 !important;
        border-radius: 12px !important;
    }

    div[data-baseweb="notification"][kind="error"] p,
    div[data-baseweb="notification"][kind="error"] div,
    div[data-baseweb="notification"][kind="error"] span {
        color: #ffe4e4 !important;
    }

    @media (max-width: 860px) {
        .main .block-container {
            padding-top: .9rem;
            padding-left: .7rem;
            padding-right: .7rem;
        }

        .title-stage {
            padding: 1rem 1rem;
            border-radius: 16px;
        }

        .title-stage h1 {
            font-size: clamp(1.25rem, 7.6vw, 2.05rem);
            line-height: 1.06;
        }

        .title-copy {
            font-size: .9rem;
        }

        .metric {
            min-height: 84px;
            padding: .72rem;
        }

        .metric .value {
            font-size: 1.02rem;
            line-height: 1.25;
        }
    }

    @keyframes fadeUp {
        from {
            opacity: 0;
            transform: translateY(8px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    @keyframes drift {
        0% {
            transform: translate3d(0, 0, 0) scale(1) rotate(0deg);
        }
        50% {
            transform: translate3d(3vw, 2vw, 0) scale(1.08) rotate(8deg);
        }
        100% {
            transform: translate3d(-2vw, 3vw, 0) scale(0.97) rotate(-7deg);
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---- sample resume ----
SAMPLE_RESUME = """John Smith
Software Engineer | Python Developer
Email: john.smith@email.com

Professional Summary
Experienced software engineer with 4 years of experience building scalable web applications
and data pipelines. Strong background in backend development and cloud infrastructure.

Technical Skills
Languages: Python, SQL, JavaScript, Java
Frameworks: Django, Flask, React.js, Node.js
Databases: PostgreSQL, MySQL, MongoDB
Tools: Git, Docker, Linux, VS Code
Cloud: AWS, Firebase
Core: REST APIs, Data Structures, Algorithms, Agile, System Design, DevOps

Work Experience
Software Engineer | TechCorp (2021 - Present)
- Developed and maintained REST APIs serving 100k daily users
- Built data pipelines using Python and SQL for analytics dashboard
- Deployed applications on AWS using Docker and Kubernetes containers
- Worked on machine learning pipelines for recommendation engine
- Collaborated in agile team using Git and Jenkins CI/CD

Junior Developer | StartupXYZ (2020 - 2021)
- Built frontend components using React.js
- Database management with PostgreSQL and MongoDB
- Cybersecurity best practices and secure coding standards
- Web development using Django and Node.js

Education
B.Tech Computer Science | XYZ University | 2020
"""

# ---- input ----
input_col, helper_col = st.columns([2.4, 1.2], gap="large")

with input_col:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    use_sample = st.checkbox("Use sample IT resume")

    uploaded_file = st.file_uploader(
        "Upload Resume File",
        type=["pdf", "docx", "doc", "txt"],
        help="Supported: PDF, DOCX, DOC, and TXT",
    )

    if uploaded_file is not None:
        try:
            ext = uploaded_file.name.lower().rsplit(".", 1)[-1]
            if ext == "pdf":
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif ext == "docx":
                extracted_text = extract_text_from_docx(uploaded_file)
            elif ext == "doc":
                extracted_text = extract_text_from_doc(uploaded_file)
            else:
                extracted_text = uploaded_file.getvalue().decode("utf-8", errors="ignore").strip()

            if extracted_text:
                st.session_state.resume_input_text = extracted_text
                st.session_state.resume_input_source = "file"
                safe_file_name = html.escape(uploaded_file.name)
                st.markdown(
                    f"<div class='upload-success'>Loaded text from <span class='upload-file-name'>{safe_file_name}</span></div>",
                    unsafe_allow_html=True,
                )
            else:
                st.warning("File uploaded, but no readable text was found.")
        except Exception as exc:
            st.error(f"Could not read file: {exc}")

    if use_sample and st.session_state.resume_input_source != "file":
        st.session_state.resume_input_text = SAMPLE_RESUME
        st.session_state.resume_input_source = "sample"
    elif not use_sample and st.session_state.resume_input_source == "sample":
        st.session_state.resume_input_text = ""
        st.session_state.resume_input_source = "manual"

    resume_input = st.text_area(
        "Paste Resume Text",
        height=320,
        placeholder="Paste your full resume text here...",
        key="resume_input_text",
    )
    analyze_clicked = st.button("Analyze Resume", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

with helper_col:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("### Input Tips")
    st.markdown("- Upload PDF or Word files, or paste text manually.")
    st.markdown("- Include summary, skills, experience, and tools.")
    st.markdown("- Keep plain text format for best extraction.")
    st.markdown("- Add role-specific keywords naturally.")
    st.markdown("<p class='small-note'>No data is persisted by this app.</p>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# ---- analyze ----
if analyze_clicked:
    if not resume_input.strip():
        st.warning("Please paste some resume text first.")
    else:
        with st.spinner("Analyzing your resume..."):
            r = analyze_resume(resume_input)

        m1, m2, m3 = st.columns(3)
        with m1:
            st.markdown(
                f"""
                <div class="metric metric-role">
                    <div class="label">Predicted Role</div>
                    <div class="value">{r['role']}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with m2:
            st.markdown(
                f"""
                <div class="metric">
                    <div class="label">Profile Match</div>
                    <div class="value">{r['match_score']}%</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with m3:
            st.markdown(
                f"""
                <div class="metric">
                    <div class="label">Role Margin</div>
                    <div class="value">{r['confidence']}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        if r['confidence'] < 15:
            st.warning("Low confidence signal: your resume may span multiple domains.")

        c1, c2 = st.columns([1.1, 1.8], gap="large")

        with c1:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown("### Top Role Candidates")
            for i, (role, score) in enumerate(r['top3'], start=1):
                st.markdown(
                    f"""
                    <div class="rank-item"><strong>{i}.</strong> {role}<br/>
                    <span class="small-note">decision score: {score}</span></div>
                    """,
                    unsafe_allow_html=True,
                )
            st.markdown("</div>", unsafe_allow_html=True)

        with c2:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown("### Match Overview")
            st.progress(int(r['match_score']) / 100)
            st.markdown(f"<p class='small-note'>Coverage score: <strong>{r['match_score']}%</strong></p>", unsafe_allow_html=True)

            if r['match_score'] >= 60:
                st.success("Strong profile alignment for this role.")
            elif r['match_score'] >= 30:
                st.warning("Moderate alignment: several important skills are still missing.")
            else:
                st.error("Low alignment: expand role-specific skills before applying.")
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("### Skill Breakdown by Priority")

        def render_skill_chips(skills, missing=False):
            if not skills:
                return "<span class='small-note'>None</span>"
            class_name = "chip miss" if missing else "chip"
            return "".join([f"<span class='{class_name}'>{s}</span>" for s in skills])

        tiers = [
            ("Core Skills", r['high_found'], r['high_missing']),
            ("Secondary Skills", r['medium_found'], r['medium_missing']),
            ("Tools", r['tools_found'], r['tools_missing']),
        ]

        for tier_name, found, missing in tiers:
            with st.expander(f"{tier_name}: {len(found)} found, {len(missing)} missing"):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Found**")
                    st.markdown(render_skill_chips(found), unsafe_allow_html=True)
                with col2:
                    st.markdown("**Missing**")
                    st.markdown(render_skill_chips(missing, missing=True), unsafe_allow_html=True)

        if r['all_missing']:
            st.markdown("### Recommendations")
            st.markdown("Focus on these skills in priority order:")

            if r['high_missing']:
                st.markdown("**High Priority: learn first**")
                for s in r['high_missing']:
                    st.markdown(f"- {s}")

            if r['medium_missing']:
                st.markdown("**Medium Priority**")
                for s in r['medium_missing']:
                    st.markdown(f"- {s}")

            if r['tools_missing']:
                st.markdown("**Tools to add**")
                for s in r['tools_missing']:
                    st.markdown(f"- {s}")

        st.caption("Note: Predictions come from a machine-learning model trained on resume text. Mixed-domain resumes can reduce certainty.")
